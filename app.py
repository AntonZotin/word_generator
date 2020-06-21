import re
import sys

import PySimpleGUI as Gui
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from datetime import datetime, date

from openpyxl import load_workbook

required_fields = {
    'name': 'Наименование компании',
    'inn': 'ИНН',
    'number': 'Номер заявки',
    'date': 'Дата заявки',
    'ispolnitel': 'Исполнитель',
    'postanovlenie': 'Постановление',
    'has_comment': 'Соответствие заявки',
    'comment': 'Комментарий',
}

numeric_fields = ['inn', 'number']

tab = '\t'
text1 = 'Юридическим отделом государственного казенного учреждения «Центр реализации программ поддержки и ' \
        'развития малого и среднего предпринимательства Республики Татарстан» (далее – Учреждение) проверена заявка '
text2 = ' (ИНН '
text3 = ') № '
text4 = ' от '
text5 = ' на предмет соответствия требованиям '
text6 = ' (далее – Порядок).'
postanovleniya = {
    'Процентная ставка': 'Порядка предоставления субсидий на возмещение части затрат, связанных с уплатой процентов '
                         'по кредитам, привлеченным в российских кредитных организациях, утвержденного постановлением '
                         'Кабинета Министров Республики Татарстан от 25.04.2020 № 327 «Об утверждении Порядка '
                         'предоставления субсидий на возмещение части затрат, связанных с уплатой процентов по '
                         'кредитам, привлеченным в российских кредитных организациях»',
    'Доставка': 'Порядка предоставления субсидий из бюджета Республики Татарстан в целях возмещения затрат субъектов '
                'малого и среднего предпринимательства, связанных с оплатой услуг (комиссии) сервисов с доставкой '
                'продуктов питания и еду, утвержденного постановлением Кабинета Министров Республики Татарстан от '
                '25.04.2020 № 326 «Об утверждении Порядка предоставления субсидий из бюджета Республики Татарстан в '
                'целях возмещения затрат субъектов малого и среднего предпринимательства, связанных с оплатой услуг ('
                'комиссии) сервисов с доставкой продуктов питания и еду» '
}
success = 'По итогам проверки замечания не выявлены.'
fail_single = 'По итогам проверки выявлено следующее замечание:'
fail_multi = 'По итогам проверки выявлены следующие замечания:'
specialists = {
    'Е.Р. Зотина': 'Главный специалист						                             Е.Р.Зотина',
    'М.Р. Галиев': 'Главный специалист						                           М.Р.Галиев',
    'Д.З. Соловьев': 'Главный специалист						                        Д.З.Соловьев'
}
footer = 'Дата: '

HAS_NO_COMMENT = 'Соответствует'
HAS_COMMENT = 'Не соответствует'
END_OF_COMMENT = '<END_OF_COMMENT>'
COMMENTS_FILE = 'templates.txt'

XLSX_FILE_PREFIX = 'Реестр заявок №'
XLSX_FILE_SUFFIX = '.xlsx'
XLSX_PP = 'C'
XLSX_NAME = 'D'
XLSX_INN = 'E'
XLSX_NUMBER = 'F'
XLSX_DATE = 'G'
XLSX_ISPOLNITEL = 'H'


def insert_values_in_row(ws, values, row):
    ws[f'{XLSX_NAME}{row}'].value = values[0]
    ws[f'{XLSX_INN}{row}'].value = values[1]
    ws[f'{XLSX_NUMBER}{row}'].value = values[2]
    ws[f'{XLSX_DATE}{row}'].value = values[3]
    ws[f'{XLSX_ISPOLNITEL}{row}'].value = values[4]


def get_str_date(ws, row):
    return str(ws[f'{XLSX_DATE}{row.row}'].value.strftime("%d.%m.%Y")) \
        if isinstance(ws[f'{XLSX_DATE}{row.row}'].value, datetime) \
           or isinstance(ws[f'{XLSX_DATE}{row.row}'].value, date) else ws[f'{XLSX_DATE}{row.row}'].value


def separate_comment(comment):
    return [re.sub('^\d+[.)\s+]+', '', c) for c in comment.strip().split('\n')]


def main_insert_and_sort_xlsx(name, inn, number, date, ispolnitel, postanovlenie):
    number_post = '327' if postanovlenie == 'Процентная ставка' else '326'
    number_prefix = 'П' if postanovlenie == 'Процентная ставка' else 'Д'
    document_name = f'{XLSX_FILE_PREFIX}{number_post}{XLSX_FILE_SUFFIX}'
    document = load_workbook(document_name)
    ws = document.active
    new_row = [name, inn, f'{number_prefix}-{number}', date, ispolnitel]
    rows = [new_row, ]
    begin_row = 1
    for row in ws[XLSX_NAME]:
        if row.row == 1:
            continue
        elif not row.value:
            if ws[f'{XLSX_PP}{row.row}'].value:
                break
            else:
                rows = [new_row, ]
                begin_row = row.row
                continue
        else:
            rows.append([
                ws[f'{XLSX_NAME}{row.row}'].value,
                ws[f'{XLSX_INN}{row.row}'].value,
                ws[f'{XLSX_NUMBER}{row.row}'].value,
                get_str_date(ws, row),
                ws[f'{XLSX_ISPOLNITEL}{row.row}'].value
            ])
    sorted_rows = sorted(rows, key=lambda r: r[2])
    for row in range(len(sorted_rows)):
        insert_values_in_row(ws, sorted_rows[row], begin_row + row + 1)
    document.save(document_name)


def main_generate_word(name, inn, number, date, ispolnitel, postanovlenie, has_comment, comment):
    number_prefix = 'П' if postanovlenie == 'Процентная ставка' else 'Д'
    fail = fail_multi if has_comment and "\n" in comment else fail_single
    comment = comment.replace("\n", "\t\n" + tab) if has_comment else ''
    paragraph1 = f'{tab}{text1}{name}{text2}{inn}{text3}{f"{number_prefix}-{number}"}{text4}{date}{text5}' \
        f'{postanovleniya[postanovlenie]}{text6}\t\n' \
        f'{tab}{fail if has_comment else success}\t\n' \
        f'{tab}{comment if has_comment else ""}\t\n'
    paragraph2 = f'{specialists[ispolnitel]}\n\n\n'
    paragraph3 = f'{footer}{datetime.today().strftime("%d.%m.%Y")}'

    document = Document()

    obj_styles = document.styles
    obj_charstyle = obj_styles.add_style('Main title', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(14)
    obj_font.name = 'Times New Roman'
    obj_charstyle = obj_styles.add_style('Middle paragraph', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(14)
    obj_font.name = 'Times New Roman'
    obj_charstyle = obj_styles.add_style('Last paragraph', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(8)
    obj_font.name = 'Times New Roman'

    t = document.add_paragraph('')
    t.add_run('ЗАКЛЮЧЕНИИЕ ЮРИДИЧЕСКОГО ОТДЕЛА', style='Main title').bold = True
    t.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    p1 = document.add_paragraph('')
    p1.add_run(paragraph1, style='Middle paragraph')
    p1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    p2 = document.add_paragraph('')
    p2.add_run(paragraph2, style='Middle paragraph').bold = True
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    p3 = document.add_paragraph('')
    p3.add_run(paragraph3, style='Last paragraph')
    p3.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    document.save("%s %s.docx" % (name.replace('\"', '\''), number[-4:]))


def main():
    with open(COMMENTS_FILE, 'r+') as t:
        comments_file = t.read().strip()
        comments_array = set(e.strip() for e in
                             filter(lambda el: el, comments_file.split(END_OF_COMMENT))) if comments_file else []
    layout = [
        [Gui.Text('Наименование компании', size=(20, 1)), Gui.InputText(size=(42, 1), key='name')],
        [Gui.Text('ИНН', size=(20, 1)), Gui.Input(size=(42, 1), key='inn', enable_events=True)],
        [Gui.Text('Номер заявки', size=(20, 1)), Gui.Input(size=(42, 1), key='number', enable_events=True)],
        [Gui.Text('Дата заявки', size=(20, 1)), Gui.Input(size=(42, 1), key='date', enable_events=True)],
        [Gui.Text('Исполнитель', size=(20, 1)), Gui.Combo([*specialists.keys()], size=(40, 1),
                                                          readonly=True, key='ispolnitel')],
        [Gui.Text('Постановление', size=(20, 1)), Gui.Combo([*postanovleniya.keys()], size=(40, 1),
                                                            readonly=True, key='postanovlenie')],
        [Gui.Text('_' * 68)],
        [Gui.Text('Соответствие заявки', size=(20, 1)),
         Gui.Radio(HAS_NO_COMMENT, default=True, group_id='1', key=HAS_NO_COMMENT),
         Gui.Radio(HAS_COMMENT, default=False, group_id='1', key=HAS_COMMENT)],
        [Gui.Text('Комментарий', size=(20, 1)), Gui.Multiline(size=(40, 10), key='comment', disabled=True)],
        [Gui.Text('Шаблон комментария', size=(20, 1)), Gui.Button('Выбрать', size=(10, 1), key='template')],
        [Gui.Text('', size=(25, 1)), Gui.Submit(button_text='Сгенерировать')]
    ]

    window = Gui.Window('Word generator', layout, finalize=True)

    window[HAS_NO_COMMENT].bind('<Button-1>', '')
    window[HAS_COMMENT].bind('<Button-1>', '')

    templates_active = False
    while True:
        event, values = window.read(timeout=100)
        print(event, values)
        if event in (None, 'Exit', 'Cancel', 'Закрыть'):
            return 0
        elif templates_active:
            template_event, template_values = template_window.Read(timeout=100)
            if template_event in (None, 'Exit', 'Cancel'):
                templates_active = False
                template_window.close()
            elif template_event == 'Select template':
                try:
                    tv = list(filter(lambda k: template_values[k], template_values.keys()))
                    window['comment'].update(re.sub('\n$', '', values['comment'], 1) + tv[0] + '\n')
                    templates_active = False
                    template_window.close()
                    window['comment'].update(disabled=False)
                except IndexError:
                    Gui.popup('Вы не выбрали шаблон', title='')
        elif not templates_active and event == 'template':
            if values[HAS_COMMENT]:
                templates_active = True
                template_layout = [
                    [Gui.Col([[Gui.Radio(ca, default=False, group_id='2', key=ca,
                                         text_color='black', background_color='white')] for ca in comments_array],
                             background_color='white', size=(600, 500), scrollable=True)],
                    [Gui.Text('', size=(30, 1)), Gui.Submit(button_text='Выбрать', key='Select template'),
                     Gui.Submit(button_text='Закрыть', key='Cancel')]]
                template_window = Gui.Window('Выбор шаблона', template_layout)
            else:
                Gui.popup('Выбери соответствие заявки: Не соответствует', title='Статус заявки')
        elif event == HAS_NO_COMMENT:
            window['comment'].update('', disabled=True)
        elif event == HAS_COMMENT:
            window['comment'].update(disabled=False)
        elif event in numeric_fields and values[event] and values[event][-1] not in '0123456789':
            window[event].update(values[event][:-1])
        elif event == 'date' and values[event]:
            if len(values[event]) == 2 or len(values[event]) == 5:
                window[event].update(values[event] + '.')
            elif len(values[event]) == 11:
                window[event].update(values[event][:-1])
        elif event == 'Сгенерировать':
            required_errors = []
            name = values['name']
            window['name'].update('')
            if not name: required_errors.append(required_fields['name'])
            inn = values['inn']
            window['inn'].update('')
            if not inn: required_errors.append(required_fields['inn'])
            number = values['number']
            window['number'].update('')
            if not number: required_errors.append(required_fields['number'])
            date = values['date']
            window['date'].update('')
            if not date: required_errors.append(required_fields['date'])
            ispolnitel = values['ispolnitel']
            window['ispolnitel'].update('')
            if not ispolnitel: required_errors.append(required_fields['ispolnitel'])
            postanovlenie = values['postanovlenie']
            window['postanovlenie'].update('')
            if not postanovlenie: required_errors.append(required_fields['postanovlenie'])
            comment = values['comment'].strip()
            window['comment'].update('')
            if not comment and values[HAS_COMMENT]: required_errors.append(required_fields['comment'])
            window[HAS_COMMENT].update(False)
            window[HAS_NO_COMMENT].update(True)

            if not required_errors:
                if comment:
                    for c in separate_comment(comment):
                        if c not in comments_array:
                            with open(COMMENTS_FILE, 'a+') as f:
                                f.write('%s%s\n' % (c, END_OF_COMMENT))
                            comments_array.add(c)
                main_generate_word(name, inn, number, date, ispolnitel, postanovlenie, values[HAS_COMMENT], comment)
                main_insert_and_sort_xlsx(name, inn, number, date, ispolnitel, postanovlenie)
            else:
                Gui.popup('Вы не ввели обязательные поля:\n%s' % ', '.join(required_errors), title='Пустые поля')


if __name__ == '__main__':
    sys.exit(main())
