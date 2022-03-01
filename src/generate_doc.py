
from docx import Document
from datetime import datetime
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

from src.strings import fail_multi, tab, text1, text2, text3, text4, text5, text6, fail_single, \
    postanovleniya, success, specialists, footer


def has_comment(data):
    return "comment" in data


def main_generate_word(data):
    fail = fail_multi if has_comment(data) and "\n" in data["comment"] else fail_single
    comment = data['comment'].replace("\n", "\t\n" + tab) if has_comment(data) else ''
    number = f"{data['postanovlenie']}-{data['number']}"
    paragraph1 = f'{tab}{text1}{data["name"]}{text2}{data["inn"]}{text3}{number}{text4}{data["request_date"]}{text5}' \
        f'{postanovleniya[data["postanovlenie"]]}{text6}\t\n' \
        f'{tab}{fail if has_comment(data) else success}\t\n' \
        f'{tab}{comment if has_comment(data) else ""}\t\n'
    paragraph2 = f'{specialists[data["ispolnitel"]]}\n\n\n'
    paragraph3 = f'{footer}{datetime.today().strftime("%d.%m.%Y")}'

    document = Document()

    obj_styles = document.styles
    obj_charstyle = obj_styles.add_style('Main title', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(14)
    obj_font.name = 'Times New Roman'
    obj_charstyle = obj_styles.add_style('Middle paragraph', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(14)
    obj_font.name = 'Times New Roman'
    obj_charstyle = obj_styles.add_style('Last paragraph', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(8)
    obj_font.name = 'Times New Roman'

    t = document.add_paragraph('')
    t.add_run('ЗАКЛЮЧЕНИИЕ ЮРИДИЧЕСКОГО ОТДЕЛА', style='Main title').bold = True
    t.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    p1 = document.add_paragraph('')
    p1.add_run(paragraph1, style='Middle paragraph')
    p1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    p2 = document.add_paragraph('')
    p2.add_run(paragraph2, style='Middle paragraph').bold = True
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    p3 = document.add_paragraph('')
    p3.add_run(paragraph3, style='Last paragraph')
    p3.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    document.save("../%s %s.docx" % (data['name'].replace('\"', '\''), number[-4:]))
