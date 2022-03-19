
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

from src.utils.strings import fail_multi, tab, text1, text2, fail_single, \
    success, specialists, footer


def has_comment(data):
    return "comment" in data


def main_generate_word(data):
    fail = fail_multi if has_comment(data) and "\n" in data["comment"] else fail_single
    comment = ''
    for c in data.get('comment', '').split('\n'):
        comment += f'{tab}{c}\n'
    number = f"{data['postanovlenie']}-{data['number']}"
    paragraph1 = f'Заявитель: {data["name"]}\n' \
        f'ИНН: {data["inn"]}\n' \
        f'№ заявки  {number}\n' \
        f'Дата: {data["request_date"]}\n'
    paragraph2 = f'{tab}{text1}{fail if has_comment(data) else success}\t\n' \
        f'{comment if has_comment(data) else ""}\n'
    paragraph3 = f'{specialists[data["ispolnitel"]]}'
    paragraph4 = f'{text2}'
    paragraph5 = f'{footer}{data["check_date"]}\n\n\n'

    document = Document()

    obj_styles = document.styles
    obj_charstyle = obj_styles.add_style('Main title', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(14)
    obj_font.name = 'Times New Roman'
    obj_charstyle = obj_styles.add_style('Middle paragraph', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(14)
    obj_font.name = 'Times New Roman'
    obj_charstyle = obj_styles.add_style('Last paragraph', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(11)
    obj_font.name = 'Times New Roman'
    obj_charstyle = obj_styles.add_style('Footer paragraph', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(8)
    obj_font.name = 'Times New Roman'

    t = document.add_paragraph('')
    t.add_run('ЗАКЛЮЧЕНИИЕ ЮРИДИЧЕСКОГО ОТДЕЛА\n\n', style='Main title')
    t.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    p1 = document.add_paragraph('')
    p1.add_run(paragraph1, style='Middle paragraph')
    p1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    p2 = document.add_paragraph('')
    p2.add_run(paragraph2, style='Middle paragraph')
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    p3 = document.add_paragraph('')
    p3.add_run(paragraph3, style='Middle paragraph')
    p3.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    p4 = document.add_paragraph('')
    p4.add_run(paragraph4, style='Last paragraph').italic = True
    p4.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    p5 = document.add_paragraph('')
    p5.add_run(paragraph5, style='Footer paragraph')
    p5.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    document.save("../%s %s.docx" % (data['name'].replace('\"', '\''), number[-4:]))
